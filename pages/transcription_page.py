from core.sounds import play_done, play_error
from core.lang import LangManager, tr, TLabel, TButton, TGroupBox, TCheckBox
"""
Module 4 : Transcription audio/vidéo

Mise en page 3 colonnes :
- Gauche  : contrôles (options, export, traduction)
- Milieu  : transcript éditable, avec timecodes et surlignage synchronisé
- Droite  : visualiseur vidéo (optionnel, si fichier vidéo)

Fonctionnalités :
- Transcription locale via faster-whisper (aucun compte requis)
- Traduction vers une autre langue via deep-translator (internet requis,
  aucun compte)
- Export TXT / DOCX / PDF (avec ou sans timecodes) / SRT / VTT
- Paramètres sous-titres : mots/sous-titre, lignes/sous-titre,
  durée minimum, espacement minimum, vitesse de lecture max (caractères/s)
- Surlignage automatique de la phrase en cours lors de la lecture vidéo
- Simple clic sur une phrase du transcript : édition du texte (correction
  d'orthographe, ponctuation...)
- Double clic sur une phrase : déplace la lecture vidéo à cet endroit
"""

import os
import re

from PySide6.QtCore import Qt, QThread, Signal, QUrl, QTimer
from PySide6.QtWidgets import (
    QWidget, QVBoxLayout, QHBoxLayout, QLabel, QPushButton, QFileDialog,
    QComboBox, QProgressBar, QTextEdit, QGroupBox, QFrame, QSpinBox,
    QMessageBox, QSplitter, QDoubleSpinBox, QCheckBox, QSlider, QSizePolicy
)
from PySide6.QtGui import QTextCursor, QTextCharFormat, QColor, QFont
from PySide6.QtMultimedia import QMediaPlayer, QAudioOutput
from PySide6.QtMultimediaWidgets import QVideoWidget

try:
    from pymediainfo import MediaInfo
except ImportError:
    MediaInfo = None


LANGUAGES = {
    "Auto-detect": None,
    "French": "fr",
    "Anglais": "en",
    "Espagnol": "es",
    "Allemand": "de",
    "Italien": "it",
    "Portugais": "pt",
    "Dutch": "nl",
    "Russe": "ru",
    "Japonais": "ja",
    "Chinois": "zh",
    "Arabe": "ar",
    "Korean": "ko",
}

TRANSLATION_LANGUAGES = {k: v for k, v in LANGUAGES.items() if v is not None}

MODEL_SIZES = {
    "Rapide (base)": "base",
    "Balanced (small)": "small",
    "Accurate (medium)": "medium",
    "Very accurate (large-v3, slow on CPU)": "large-v3",
}

EXPORT_FORMATS = ["TXT", "DOCX (Word)", "PDF", "SRT", "VTT"]
SUBTITLE_FORMATS = {"SRT", "VTT"}

HIGHLIGHT_BG_COLOR = QColor("#ededed")
HIGHLIGHT_TEXT_COLOR = QColor("#e8542e")
_TC_PREFIX_PATTERN = re.compile(r"^\[\d{2}:\d{2}\]\s*")


def _get_duration_seconds(path):
    # 1) pymediainfo si disponible
    if MediaInfo is not None:
        try:
            info = MediaInfo.parse(path)
            for track in info.tracks:
                if track.track_type == "General":
                    duration_ms = getattr(track, "duration", None)
                    if duration_ms not in (None, ""):
                        return float(duration_ms) / 1000
        except Exception:
            pass
    # 2) Secours ffprobe (indispensable sur macOS sans pymediainfo)
    try:
        from core.ffprobe_reader import read_metadata
        meta = read_metadata(path)
        if meta and meta.get("general", {}).get("duration_s") is not None:
            return float(meta["general"]["duration_s"])
    except Exception:
        pass
    return None


def _format_srt_ts(seconds):
    ms = int(round(seconds * 1000))
    h, ms = divmod(ms, 3600000)
    m, ms = divmod(ms, 60000)
    s, ms = divmod(ms, 1000)
    return f"{h:02d}:{m:02d}:{s:02d},{ms:03d}"


def _format_vtt_ts(seconds):
    ms = int(round(seconds * 1000))
    h, ms = divmod(ms, 3600000)
    m, ms = divmod(ms, 60000)
    s, ms = divmod(ms, 1000)
    return f"{h:02d}:{m:02d}:{s:02d}.{ms:03d}"


def _format_display_tc(seconds):
    """Timecode court pour l'affichage dans le transcript : MM:SS"""
    m = int(seconds // 60)
    s = int(seconds % 60)
    return f"[{m:02d}:{s:02d}]"


def _wrap_into_lines(text, num_lines):
    words = text.split()
    if num_lines <= 1 or len(words) <= 1:
        return text
    per_line = max(1, round(len(words) / num_lines))
    return "\n".join(
        " ".join(words[i:i + per_line])
        for i in range(0, len(words), per_line)
    )


def _apply_subtitle_timing_rules(cues, min_dur, min_gap, max_chars_per_sec):
    """Ajuste les durées d'affichage des cues selon les règles pro :
    - durée minimale d'affichage (norme courante : 0.8 s)
    - espacement minimum entre deux cues, pour bien les distinguer (norme : 0.08 s)
    - durée minimale basée sur la vitesse de lecture, pour laisser le temps
      de lire un sous-titre long (norme streaming/broadcast : ~17 car/s)
    """
    result = []
    for cue in cues:
        start = cue["start"]
        end = cue["end"]
        text = cue["text"]

        if max_chars_per_sec > 0:
            min_by_speed = len(text) / max_chars_per_sec
            end = max(end, start + min_by_speed)

        if min_dur > 0:
            end = max(end, start + min_dur)

        result.append({"start": start, "end": end, "text": text})

    for i in range(1, len(result)):
        prev_end = result[i - 1]["end"]
        curr_start = result[i]["start"]
        if prev_end + min_gap > curr_start:
            result[i - 1]["end"] = max(result[i - 1]["start"] + 0.1,
                                       curr_start - min_gap)
    return result


class AudioDropZone(QFrame):
    file_dropped = Signal(str)

    def __init__(self, parent=None):
        super().__init__(parent)
        self.setAcceptDrops(True)
        self.setFrameShape(QFrame.StyledPanel)
        self.setMinimumHeight(90)
        self.setStyleSheet(
            "QFrame { border: 2px solid #e8542e; border-radius: 8px; background: #1e2235; }"
        )
        layout = QVBoxLayout(self)
        self.label = QLabel("Drop an audio or video file here\n(or click \"Browse\" below)")
        self.label.setAlignment(Qt.AlignCenter)
        self.label.setStyleSheet("border: none; color: #666666;")
        layout.addWidget(self.label)

    def dragEnterEvent(self, event):
        if event.mimeData().hasUrls():
            event.acceptProposedAction()
        self.setStyleSheet("QFrame { border: 2px solid #ff6b45; border-radius: 8px; background: #252a42; }")
    def dragLeaveEvent(self, event):
        self.setStyleSheet("QFrame { border: 2px solid #e8542e; border-radius: 8px; background: #1e2235; }")

    def dropEvent(self, event):
        self.setStyleSheet("QFrame { border: 2px solid #e8542e; border-radius: 8px; background: #1e2235; }")
        urls = event.mimeData().urls()
        if urls:
            path = urls[0].toLocalFile()
            if path:
                self.file_dropped.emit(path)


class TranscriptEditor(QTextEdit):
    """QTextEdit spécialisé pour le transcript :
    - Simple clic : place le curseur normalement, pour éditer le texte
      (correction d'orthographe, ponctuation...), comportement standard.
    - Double clic : au lieu de sélectionner un mot (comportement Qt par
      défaut), déplace la lecture vidéo à l'instant correspondant à la
      phrase cliquée.
    - Touche Entrée désactivée : on ne crée jamais de nouvelle ligne, pour
      garder la correspondance 1 ligne = 1 segment transcrit (nécessaire
      pour le surlignage synchronisé et l'export).
    """
    segment_double_clicked = Signal(int)

    def mouseDoubleClickEvent(self, event):
        cursor = self.cursorForPosition(event.pos())
        self.segment_double_clicked.emit(cursor.blockNumber())
        # Ne pas appeler super() : on évite la sélection de mot par défaut.

    def keyPressEvent(self, event):
        if event.key() in (Qt.Key_Return, Qt.Key_Enter):
            return
        super().keyPressEvent(event)


class TranscribeWorker(QThread):
    log_message = Signal(str)
    progress_value = Signal(int)
    segment_ready = Signal(dict)
    finished_all = Signal(bool, list)

    def __init__(self, audio_path, model_size, language_code, parent=None):
        super().__init__(parent)
        self.audio_path = audio_path
        self.model_size = model_size
        self.language_code = language_code
        self._stop_requested = False

    def stop(self):
        self._stop_requested = True

    def run(self):
        try:
            from faster_whisper import WhisperModel
        except ImportError:
            self.log_message.emit("ERREUR : faster-whisper n'est pas installé. Lance : pip install -r requirements.txt")
            self.finished_all.emit(False, [])
            return

        try:
            model = WhisperModel(self.model_size, device="cpu", compute_type="int8")
        except Exception as e:
            self.log_message.emit(f"ERROR loading model: {e}")
            self.finished_all.emit(False, [])
            return

        duration = _get_duration_seconds(self.audio_path)
        self.log_message.emit("Transcription en cours...")

        try:
            segments_gen, info = model.transcribe(
                self.audio_path,
                language=self.language_code,
                word_timestamps=True,
            )
        except Exception as e:
            self.log_message.emit(f"ERREUR lors de la transcription : {e}")
            self.finished_all.emit(False, [])
            return

        segments_data = []
        for seg in segments_gen:
            if self._stop_requested:
                self.log_message.emit("Cancelled.")
                self.finished_all.emit(False, segments_data)
                return

            words = [{"start": w.start, "end": w.end, "word": w.word} for w in (seg.words or [])]
            seg_dict = {"start": seg.start, "end": seg.end, "text": seg.text.strip(), "words": words}
            segments_data.append(seg_dict)
            self.segment_ready.emit(seg_dict)

            if duration:
                pct = max(0, min(100, int(seg.end / duration * 100)))
                self.progress_value.emit(pct)

        self.progress_value.emit(100)
        self.log_message.emit("✔ Transcription complete.")
        self.finished_all.emit(True, segments_data)


class TranslateWorker(QThread):
    log_message = Signal(str)
    finished_translation = Signal(bool, list)

    def __init__(self, segments_data, target_lang_code, parent=None):
        super().__init__(parent)
        self.segments_data = segments_data
        self.target_lang_code = target_lang_code

    def run(self):
        try:
            from deep_translator import GoogleTranslator
        except ImportError:
            self.log_message.emit("ERREUR : deep-translator n'est pas installé. Lance : pip install -r requirements.txt")
            self.finished_translation.emit(False, [])
            return

        self.log_message.emit("Traduction en cours...")
        try:
            translator = GoogleTranslator(source="auto", target=self.target_lang_code)
            translated = []
            for seg in self.segments_data:
                translated_text = translator.translate(seg["text"])
                translated.append({**seg, "text": translated_text or seg["text"]})
            self.log_message.emit("✔ Translation complete.")
            self.finished_translation.emit(True, translated)
        except Exception as e:
            self.log_message.emit(f"ERREUR lors de la traduction : {e}")
            self.finished_translation.emit(False, [])


class TranscriptionPage(QWidget):
    back_requested = Signal()

    def __init__(self, parent=None):
        super().__init__(parent)
        self.worker = None
        self.translate_worker = None
        self.audio_path = None
        self.segments_data = []
        self._last_highlighted_block = None
        self._highlight_timer = QTimer(self)
        self._highlight_timer.setInterval(150)
        self._highlight_timer.timeout.connect(self._sync_highlight)
        self._build_ui()

    def _build_ui(self):
        outer = QVBoxLayout(self)
        outer.setContentsMargins(24, 24, 24, 16)
        outer.setSpacing(8)

        top_bar = QHBoxLayout()
        self._t_back = QPushButton(tr("back")); back_btn = self._t_back
        back_btn.clicked.connect(self.back_requested.emit)
        top_bar.addWidget(back_btn)
        top_bar.addStretch()
        outer.addLayout(top_bar)

        self._t_title = QLabel(tr("tr_title")); title = self._t_title
        title.setStyleSheet("font-size: 18px; font-weight: bold;")
        outer.addWidget(title)

        # ══ [ réglages pleine hauteur ] | [ vidéo au-dessus / transcript ] ══
        main_hsplit = QSplitter(Qt.Horizontal)
        right_vsplit = QSplitter(Qt.Vertical)

        # ═══ Vidéo (en haut à droite, centrée 16:9) ═══
        video_container = QWidget()
        video_outer = QVBoxLayout(video_container)
        video_outer.setContentsMargins(0, 0, 0, 6)
        video_outer.addWidget(QLabel("Preview:"))

        video_center_row = QHBoxLayout()
        video_center_row.addStretch()
        self.video_widget = QVideoWidget()
        self.video_widget.setMinimumSize(480, 270)   # 16:9, s'agrandit
        self.video_widget.setStyleSheet("background-color: black;")
        video_center_row.addWidget(self.video_widget)
        video_center_row.addStretch()
        video_outer.addLayout(video_center_row, stretch=1)

        player_controls = QHBoxLayout()
        self.play_btn = QPushButton("▶")
        self.play_btn.setFixedWidth(32)
        self.play_btn.clicked.connect(self._toggle_play)
        player_controls.addWidget(self.play_btn)
        self.stop_btn = QPushButton("■")
        self.stop_btn.setFixedWidth(32)
        self.stop_btn.clicked.connect(self._stop_player)
        player_controls.addWidget(self.stop_btn)
        self.seek_slider = QSlider(Qt.Horizontal)
        self.seek_slider.sliderMoved.connect(self._on_seek)
        player_controls.addWidget(self.seek_slider)
        video_outer.addLayout(player_controls)

        right_vsplit.addWidget(video_container)

        # ═══ GAUCHE : réglages pleine hauteur, dans une zone défilable ═══
        from PySide6.QtWidgets import QScrollArea
        left_scroll = QScrollArea()
        left_scroll.setWidgetResizable(True)
        left_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
        left_scroll.setMinimumWidth(380)
        left_scroll.setMaximumWidth(440)
        left_scroll.setFrameShape(QFrame.NoFrame)

        left_widget = QWidget()
        layout = QVBoxLayout(left_widget)
        layout.setContentsMargins(0, 0, 8, 0)
        layout.setSpacing(8)

        self.drop_zone = AudioDropZone()
        self.drop_zone.file_dropped.connect(self._load_file)
        layout.addWidget(self.drop_zone)

        browse_layout = QHBoxLayout()
        browse_btn = QPushButton("Browse...")
        browse_btn.clicked.connect(self._browse_file)
        browse_layout.addWidget(browse_btn)
        self.file_label = QLabel("No file.")
        self.file_label.setStyleSheet("color: #555555;")
        browse_layout.addWidget(self.file_label)
        browse_layout.addStretch()
        layout.addLayout(browse_layout)

        options_group = QGroupBox("Options")
        options_layout = QHBoxLayout(options_group)
        lang_box = QVBoxLayout()
        lang_box.addWidget(TLabel("language"))
        self.lang_combo = QComboBox()
        self.lang_combo.addItems(list(LANGUAGES.keys()))
        lang_box.addWidget(self.lang_combo)
        options_layout.addLayout(lang_box)
        model_box = QVBoxLayout()
        model_box.addWidget(TLabel("model"))
        self.model_combo = QComboBox()
        self.model_combo.addItems(list(MODEL_SIZES.keys()))
        model_box.addWidget(self.model_combo)
        options_layout.addLayout(model_box)
        layout.addWidget(options_group)

        action_layout = QHBoxLayout()
        self.transcribe_btn = TButton("tr_btn")
        self.transcribe_btn.setMinimumHeight(34)
        self.transcribe_btn.clicked.connect(self._start_transcribe)
        action_layout.addWidget(self.transcribe_btn)
        self.cancel_btn = TButton("cancel")
        self.cancel_btn.setEnabled(False)
        self.cancel_btn.clicked.connect(self._cancel_transcribe)
        action_layout.addWidget(self.cancel_btn)
        layout.addLayout(action_layout)

        self.progress_bar = QProgressBar()
        layout.addWidget(self.progress_bar)

        layout.addWidget(TLabel("journal"))
        self.log_output = QTextEdit()
        self.log_output.setReadOnly(True)
        self.log_output.setFixedHeight(60)
        layout.addWidget(self.log_output)

        # Traduction
        translate_group = TGroupBox("tr_translate")
        translate_layout = QHBoxLayout(translate_group)
        translate_layout.addWidget(TLabel("translate_to"))
        self.translate_lang_combo = QComboBox()
        self.translate_lang_combo.addItems(list(TRANSLATION_LANGUAGES.keys()))
        translate_layout.addWidget(self.translate_lang_combo)
        self.translate_btn = TButton("translate_btn")
        self.translate_btn.clicked.connect(self._start_translate)
        translate_layout.addWidget(self.translate_btn)
        layout.addWidget(translate_group)

        # Export
        export_group = TGroupBox("export_group")
        export_layout = QVBoxLayout(export_group)
        fmt_row = QHBoxLayout()
        fmt_row.addWidget(TLabel("format"))
        self.export_format_combo = QComboBox()
        self.export_format_combo.addItems(EXPORT_FORMATS)
        self.export_format_combo.currentTextChanged.connect(self._on_export_format_changed)
        fmt_row.addWidget(self.export_format_combo)
        export_layout.addLayout(fmt_row)

        self.subtitle_params_group = TGroupBox("subtitle_params")
        sp_outer_layout = QVBoxLayout(self.subtitle_params_group)
        sp_layout = QHBoxLayout()

        left_sp = QVBoxLayout()
        left_sp.addWidget(TLabel("words_per_sub"))
        self.words_per_cue_spin = QSpinBox()
        self.words_per_cue_spin.setRange(1, 20)
        self.words_per_cue_spin.setValue(8)
        left_sp.addWidget(self.words_per_cue_spin)
        left_sp.addWidget(TLabel("lines_per_sub"))
        self.lines_per_cue_spin = QSpinBox()
        self.lines_per_cue_spin.setRange(1, 3)
        self.lines_per_cue_spin.setValue(1)
        left_sp.addWidget(self.lines_per_cue_spin)
        sp_layout.addLayout(left_sp)

        right_sp = QVBoxLayout()
        right_sp.addWidget(TLabel("min_dur"))
        self.min_dur_spin = QDoubleSpinBox()
        self.min_dur_spin.setRange(0, 5)
        self.min_dur_spin.setSingleStep(0.1)
        self.min_dur_spin.setValue(0.8)
        right_sp.addWidget(self.min_dur_spin)
        right_sp.addWidget(TLabel("min_gap"))
        self.min_gap_spin = QDoubleSpinBox()
        self.min_gap_spin.setRange(0, 1)
        self.min_gap_spin.setSingleStep(0.01)
        self.min_gap_spin.setValue(0.08)
        right_sp.addWidget(self.min_gap_spin)
        right_sp.addWidget(TLabel("max_cps"))
        self.max_cps_spin = QDoubleSpinBox()
        self.max_cps_spin.setRange(0, 50)
        self.max_cps_spin.setSingleStep(1)
        self.max_cps_spin.setValue(17)
        right_sp.addWidget(self.max_cps_spin)
        sp_layout.addLayout(right_sp)
        sp_outer_layout.addLayout(sp_layout)

        export_layout.addWidget(self.subtitle_params_group)
        self.subtitle_params_group.setVisible(False)

        export_btn = TButton("export")
        export_btn.clicked.connect(self._export)
        export_layout.addWidget(export_btn)
        layout.addWidget(export_group)
        layout.addStretch()

        # Colonne réglages dans le scroll, à gauche du split principal
        left_scroll.setWidget(left_widget)
        main_hsplit.addWidget(left_scroll)

        # ── Transcript (en bas à droite) ───────────────────────────────
        middle_widget = QWidget()
        middle_layout = QVBoxLayout(middle_widget)
        middle_layout.setContentsMargins(0, 6, 0, 0)
        tc_row = QHBoxLayout()
        tc_row.addWidget(TLabel("transcript_lbl"))
        self.show_tc_check = TCheckBox("timecodes")
        self.show_tc_check.setChecked(True)
        self.show_tc_check.toggled.connect(self._refresh_preview)
        tc_row.addWidget(self.show_tc_check)
        tc_row.addStretch()
        middle_layout.addLayout(tc_row)
        hint_label = QLabel("Single click: edit text. Double-click: jump to that moment in the video.")
        hint_label.setStyleSheet("color: #888888; font-size: 11px;")
        middle_layout.addWidget(hint_label)
        self.transcript_view = TranscriptEditor()
        self.transcript_view.setReadOnly(True)
        self.transcript_view.setFont(QFont("Courier", 10))
        self.transcript_view.segment_double_clicked.connect(self._on_segment_double_clicked)
        middle_layout.addWidget(self.transcript_view)

        # Transcript sous la vidéo dans le splitter vertical de droite
        right_vsplit.addWidget(middle_widget)
        right_vsplit.setStretchFactor(0, 3)   # vidéo
        right_vsplit.setStretchFactor(1, 4)   # transcript
        right_vsplit.setSizes([340, 420])

        # Split principal : réglages (gauche) | vidéo+transcript (droite)
        main_hsplit.addWidget(right_vsplit)
        main_hsplit.setStretchFactor(0, 0)
        main_hsplit.setStretchFactor(1, 1)
        main_hsplit.setSizes([400, 700])
        outer.addWidget(main_hsplit, stretch=1)

        self.player = QMediaPlayer(self)
        self.audio_output = QAudioOutput(self)
        self.player.setAudioOutput(self.audio_output)
        self.player.setVideoOutput(self.video_widget)
        self.audio_output.setVolume(0.8)
        self.player.durationChanged.connect(self._on_player_duration_changed)
        self.player.positionChanged.connect(self._on_player_position_changed)
        self.player.playbackStateChanged.connect(self._on_player_state_changed)

    # ── Chargement de fichier ──────────────────────────────────────────

    def _browse_file(self):
        path, _ = QFileDialog.getOpenFileName(
            self, "Select an audio or video file", "",
            "Audio and video (*.mp3 *.wav *.m4a *.flac *.mp4 *.mov *.mkv);;All (*.*)"
        )
        if path:
            self._load_file(path)

    def _load_file(self, path):
        if not os.path.isfile(path):
            return
        self.audio_path = path
        self.file_label.setText(os.path.basename(path))
        self.transcript_view.clear()
        self.transcript_view.setReadOnly(True)
        self.segments_data = []
        self._last_highlighted_block = None
        self.player.setSource(QUrl.fromLocalFile(path))

    # ── Transcription ──────────────────────────────────────────────────

    def _start_transcribe(self):
        if not self.audio_path:
            QMessageBox.warning(self, "No file", "Drop or select a file first.")
            return
        self.transcribe_btn.setEnabled(False)
        self.cancel_btn.setEnabled(True)
        self.log_output.clear()
        self.transcript_view.clear()
        self.transcript_view.setReadOnly(True)
        self.progress_bar.setValue(0)
        self.segments_data = []
        self._last_highlighted_block = None

        self.worker = TranscribeWorker(
            audio_path=self.audio_path,
            model_size=MODEL_SIZES[self.model_combo.currentText()],
            language_code=LANGUAGES[self.lang_combo.currentText()],
        )
        self.worker.log_message.connect(self.log_output.append)
        self.worker.progress_value.connect(self.progress_bar.setValue)
        self.worker.segment_ready.connect(self._on_segment_ready)
        self.worker.finished_all.connect(self._on_transcribe_finished)
        self.worker.start()

    def _cancel_transcribe(self):
        if self.worker:
            self.worker.stop()
            self.log_output.append("Cancellation requested...")

    def _on_segment_ready(self, seg_dict):
        self.segments_data.append(seg_dict)
        cursor = self.transcript_view.textCursor()
        cursor.movePosition(QTextCursor.End)
        if self.show_tc_check.isChecked():
            cursor.insertText(f"{_format_display_tc(seg_dict['start'])} ")
        cursor.insertText(seg_dict["text"] + "\n")
        self.transcript_view.setTextCursor(cursor)

    def _on_transcribe_finished(self, success, segments_data):
        if success:
            play_done()
        else:
            play_error()
        self.transcribe_btn.setEnabled(True)
        self.cancel_btn.setEnabled(False)
        self.segments_data = segments_data
        if success:
            self._refresh_preview()
            self.transcript_view.setReadOnly(False)
            self._highlight_timer.start()

    # ── Traduction ─────────────────────────────────────────────────────

    def _start_translate(self):
        if not self.segments_data:
            QMessageBox.warning(self, "Nothing to translate", "Transcris d'abord un fichier.")
            return
        self._sync_edits_from_view()
        target_code = TRANSLATION_LANGUAGES[self.translate_lang_combo.currentText()]
        self.translate_btn.setEnabled(False)
        self.log_output.append("Traduction en cours...")

        self.translate_worker = TranslateWorker(
            segments_data=list(self.segments_data),
            target_lang_code=target_code,
        )
        self.translate_worker.log_message.connect(self.log_output.append)
        self.translate_worker.finished_translation.connect(self._on_translation_finished)
        self.translate_worker.start()

    def _on_translation_finished(self, success, translated_segments):
        self.translate_btn.setEnabled(True)
        if success and translated_segments:
            self.segments_data = translated_segments
            self._refresh_preview()
            self.transcript_view.setReadOnly(False)

    # ── Aperçu transcript ──────────────────────────────────────────────

    def _refresh_preview(self):
        show_tc = self.show_tc_check.isChecked()
        self._last_highlighted_block = None
        self.transcript_view.clear()
        for seg in self.segments_data:
            tc = _format_display_tc(seg["start"]) + " " if show_tc else ""
            self.transcript_view.append(tc + seg["text"])

    def _sync_edits_from_view(self):
        """Récupère le texte éventuellement édité à la main dans le
        transcript et le reporte dans segments_data (pour que l'export
        utilise bien le texte corrigé)."""
        doc = self.transcript_view.document()
        show_tc = self.show_tc_check.isChecked()
        for i, seg in enumerate(self.segments_data):
            block = doc.findBlockByNumber(i)
            if not block.isValid():
                continue
            text = block.text()
            if show_tc:
                text = _TC_PREFIX_PATTERN.sub("", text)
            new_text = text.strip()
            if new_text and new_text != seg["text"]:
                seg["text"] = new_text
                self._resync_words_if_edited(seg)

    def _resync_words_if_edited(self, seg):
        """Si le nombre de mots a changé suite à une édition manuelle, les
        timestamps mot-par-mot d'origine (issus de Whisper) ne correspondent
        plus : on les recalcule en répartissant la durée du segment de façon
        régulière sur les nouveaux mots (moins précis que Whisper, mais
        cohérent avec le texte édité)."""
        words_text = seg["text"].split()
        if len(words_text) == len(seg["words"]):
            return
        total_dur = max(seg["end"] - seg["start"], 0.01)
        n = len(words_text)
        if n == 0:
            seg["words"] = []
            return
        step = total_dur / n
        seg["words"] = [
            {"start": seg["start"] + i * step, "end": seg["start"] + (i + 1) * step,
             "word": (" " if i > 0 else "") + w}
            for i, w in enumerate(words_text)
        ]

    # ── Lecteur vidéo ─────────────────────────────────────────────────

    def _toggle_play(self):
        if self.player.playbackState() == QMediaPlayer.PlaybackState.PlayingState:
            self.player.pause()
        else:
            self.player.play()

    def _stop_player(self):
        self.player.stop()

    def _on_player_state_changed(self, state):
        self.play_btn.setText(
            "⏸" if state == QMediaPlayer.PlaybackState.PlayingState else "▶"
        )

    def _on_player_duration_changed(self, ms):
        self.seek_slider.setRange(0, ms)

    def _on_player_position_changed(self, ms):
        if not self.seek_slider.isSliderDown():
            self.seek_slider.setValue(ms)

    def _on_seek(self, ms):
        self.player.setPosition(ms)

    def _on_segment_double_clicked(self, block_number):
        if 0 <= block_number < len(self.segments_data):
            seg = self.segments_data[block_number]
            self.player.setPosition(int(seg["start"] * 1000))
            self._highlight_block(block_number)

    # ── Surlignage synchronisé ─────────────────────────────────────────

    def _highlight_block(self, index):
        """Surligne le segment d'index donné (fond gris clair + texte en
        couleur accent), efface le surlignage précédent, et fait défiler
        la vue pour qu'il reste visible. Utilisé à la fois pendant la
        lecture vidéo et lors d'un double-clic sur une phrase."""
        if index is None or index == self._last_highlighted_block:
            return
        if not (0 <= index < len(self.segments_data)):
            return

        doc = self.transcript_view.document()
        block = doc.findBlockByNumber(index)
        if not block.isValid():
            return

        # Effacer le surlignage du bloc précédent uniquement (pas tout le
        # document, pour ne jamais affecter le texte qui n'est pas concerné).
        if self._last_highlighted_block is not None:
            prev_block = doc.findBlockByNumber(self._last_highlighted_block)
            if prev_block.isValid():
                reset_cursor = QTextCursor(prev_block)
                reset_cursor.movePosition(QTextCursor.StartOfBlock)
                reset_cursor.movePosition(QTextCursor.EndOfBlock, QTextCursor.KeepAnchor)
                reset_cursor.setCharFormat(QTextCharFormat())

        # Surligner le bloc entier (et non la "ligne visuelle" : une phrase
        # qui retourne à la ligne doit être surlignée en entier).
        highlight_fmt = QTextCharFormat()
        highlight_fmt.setBackground(HIGHLIGHT_BG_COLOR)
        highlight_fmt.setForeground(HIGHLIGHT_TEXT_COLOR)
        hl_cursor = QTextCursor(block)
        hl_cursor.movePosition(QTextCursor.StartOfBlock)
        hl_cursor.movePosition(QTextCursor.EndOfBlock, QTextCursor.KeepAnchor)
        hl_cursor.setCharFormat(highlight_fmt)
        self._last_highlighted_block = index

        # Déplacement du curseur pour le défilement automatique, SANS créer
        # de sélection visible (sinon tout le texte apparaît "sélectionné").
        scroll_cursor = QTextCursor(block)
        scroll_cursor.clearSelection()
        self.transcript_view.setTextCursor(scroll_cursor)
        self.transcript_view.ensureCursorVisible()

    def _sync_highlight(self):
        if not self.segments_data:
            return
        if self.player.playbackState() != QMediaPlayer.PlaybackState.PlayingState:
            return

        current_sec = self.player.position() / 1000
        current_idx = None
        for i, seg in enumerate(self.segments_data):
            if seg["start"] <= current_sec <= seg["end"]:
                current_idx = i
                break
        if current_idx is None:
            # Entre deux segments (silence) : on garde le dernier segment
            # passé plutôt que de ne rien surligner (évite l'effet de saut).
            for i, seg in enumerate(self.segments_data):
                if seg["start"] <= current_sec:
                    current_idx = i
                else:
                    break

        self._highlight_block(current_idx)

    # ── Export ────────────────────────────────────────────────────────

    def _on_export_format_changed(self, fmt):
        self.subtitle_params_group.setVisible(fmt in SUBTITLE_FORMATS)

    def _export(self):
        if not self.segments_data:
            QMessageBox.warning(self, "Nothing to export", "Transcris d'abord un fichier.")
            return
        self._sync_edits_from_view()
        fmt = self.export_format_combo.currentText()
        filters = {
            "TXT": "Fichier texte (*.txt)",
            "DOCX (Word)": "Document Word (*.docx)",
            "PDF": "Document PDF (*.pdf)",
            "SRT": "Sous-titres SRT (*.srt)",
            "VTT": "Sous-titres VTT (*.vtt)",
        }
        path, _ = QFileDialog.getSaveFileName(self, "Exporter", "", filters[fmt])
        if not path:
            return
        try:
            if fmt == "TXT":
                self._export_txt(path)
            elif fmt == "DOCX (Word)":
                self._export_docx(path)
            elif fmt == "PDF":
                self._export_pdf(path)
            elif fmt == "SRT":
                self._export_srt(path)
            else:
                self._export_vtt(path)
            QMessageBox.information(self, "Export successful", f"File saved:\n{path}")
        except ImportError as e:
            QMessageBox.critical(self, "Missing dependency",
                f"{e}\nLance : pip install -r requirements.txt")
        except Exception as e:
            QMessageBox.critical(self, "Erreur d'export", str(e))

    def _build_plain_text(self, with_timecodes=False):
        lines = []
        for seg in self.segments_data:
            tc = _format_display_tc(seg["start"]) + " " if with_timecodes else ""
            lines.append(tc + seg["text"])
        return "\n\n".join(lines)

    def _build_cues(self):
        words_per_cue = self.words_per_cue_spin.value()
        lines_per_cue = self.lines_per_cue_spin.value()
        min_dur = self.min_dur_spin.value()
        min_gap = self.min_gap_spin.value()
        max_cps = self.max_cps_spin.value()

        all_words = []
        for seg in self.segments_data:
            all_words.extend(seg["words"])

        if not all_words:
            raw_cues = [{"start": s["start"], "end": s["end"], "text": s["text"]}
                        for s in self.segments_data]
        else:
            raw_cues = []
            for i in range(0, len(all_words), words_per_cue):
                chunk = all_words[i:i + words_per_cue]
                text = "".join(w["word"] for w in chunk).strip()
                raw_cues.append({"start": chunk[0]["start"], "end": chunk[-1]["end"], "text": text})

        adjusted = _apply_subtitle_timing_rules(raw_cues, min_dur, min_gap, max_cps)
        for cue in adjusted:
            cue["text"] = _wrap_into_lines(cue["text"], lines_per_cue)
        return adjusted

    def _export_txt(self, path):
        with open(path, "w", encoding="utf-8") as f:
            f.write(self._build_plain_text(with_timecodes=self.show_tc_check.isChecked()))

    def _export_docx(self, path):
        from docx import Document
        doc = Document()
        doc.add_heading("Transcript", level=1)
        for seg in self.segments_data:
            tc = _format_display_tc(seg["start"]) + " " if self.show_tc_check.isChecked() else ""
            doc.add_paragraph(tc + seg["text"])
        doc.save(path)

    def _export_pdf(self, path):
        from reportlab.lib.pagesizes import A4
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet
        doc = SimpleDocTemplate(path, pagesize=A4)
        styles = getSampleStyleSheet()
        elements = [Paragraph("Transcript", styles["Title"]), Spacer(1, 12)]
        for seg in self.segments_data:
            tc = _format_display_tc(seg["start"]) + " " if self.show_tc_check.isChecked() else ""
            elements.append(Paragraph(tc + seg["text"], styles["BodyText"]))
            elements.append(Spacer(1, 6))
        doc.build(elements)

    def _export_srt(self, path):
        cues = self._build_cues()
        lines = []
        for idx, cue in enumerate(cues, start=1):
            lines += [str(idx),
                      f"{_format_srt_ts(cue['start'])} --> {_format_srt_ts(cue['end'])}",
                      cue["text"], ""]
        with open(path, "w", encoding="utf-8") as f:
            f.write("\n".join(lines))

    def _export_vtt(self, path):
        cues = self._build_cues()
        lines = ["WEBVTT", ""]
        for cue in cues:
            lines += [f"{_format_vtt_ts(cue['start'])} --> {_format_vtt_ts(cue['end'])}",
                      cue["text"], ""]
        with open(path, "w", encoding="utf-8") as f:
            f.write("\n".join(lines))

    def set_lang(self, lang):
        from core.lang import tr, LangManager
        LangManager.get().current = lang
        if hasattr(self, '_t_title'):     self._t_title.setText(tr("tr_title"))
        if hasattr(self, '_t_back'):      self._t_back.setText(tr("back"))
        if hasattr(self, '_t_tr_btn'):    self._t_tr_btn.setText(tr("tr_btn"))
        if hasattr(self, '_t_cancel'):    self._t_cancel.setText(tr("cancel"))

