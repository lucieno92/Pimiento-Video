from core.sounds import play_done, play_error
from core.lang import LangManager, tr, TLabel, TButton, TGroupBox, TCheckBox
"""
Module 5 : Feuille de droits musicaux (Cue Sheet)

Importer un EDL (format CMX3600), sélectionner les pistes audio contenant
la musique, générer un tableau (nom de la musique, TC IN, TC OUT, durée)
exportable en PDF, Word ou TXT.

L'EDL est exporté nativement par Avid Media Composer, Premiere Pro,
DaVinci Resolve et Final Cut Pro — tous les logiciels de montage courants.
"""

import os

from PySide6.QtCore import Qt, Signal
from PySide6.QtWidgets import (
    QWidget, QVBoxLayout, QHBoxLayout, QLabel, QPushButton, QFileDialog,
    QGroupBox, QFrame, QCheckBox, QLineEdit, QTableWidget, QTableWidgetItem,
    QMessageBox, QHeaderView, QComboBox
)

from core.edl_parser import parse_edl, timecode_to_frames, frames_to_timecode, EdlParseError


COLUMNS = ["Evt", "Music / clip name", "Track", "TC IN", "TC OUT", "Duration"]

FRAME_RATES = {
    "23.976": 23.976,
    "24": 24,
    "25 (PAL)": 25,
    "29.97 (NTSC drop-frame)": 29.97,
    "30": 30,
    "50": 50,
    "60": 60,
}


class EdlDropZone(QFrame):
    file_dropped = Signal(str)

    def __init__(self, parent=None):
        super().__init__(parent)
        self.setAcceptDrops(True)
        self.setFrameShape(QFrame.StyledPanel)
        self.setMinimumHeight(90)
        self.setStyleSheet(
            "QFrame { border: 2px solid #e8542e; border-radius: 8px; background: #1e2235; }"
        )
        layout = QVBoxLayout(self)
        self.label = QLabel(
            "Drop an EDL file here\n"
            "(Avid / Premiere Pro / DaVinci Resolve / Final Cut Pro — format CMX3600)"
        )
        self.label.setAlignment(Qt.AlignCenter)
        self.label.setStyleSheet("border: none; color: #666666;")
        self.label.setWordWrap(True)
        from PySide6.QtWidgets import QSizePolicy as _SP
        self.label.setSizePolicy(_SP.Ignored, _SP.Preferred)
        layout.addWidget(self.label)

    def show_filename(self, path):
        import os as _os
        from PySide6.QtGui import QFontMetrics
        name = _os.path.basename(path)
        self.label.setStyleSheet("border: none; color: #e8eaf0; font-weight: 600;")
        avail = max(self.width() - 24, 180)
        self.label.setText(QFontMetrics(self.label.font()).elidedText(
            name, Qt.ElideMiddle, avail))
        self.label.setToolTip(name)

    def dragEnterEvent(self, event):
        if event.mimeData().hasUrls():
            event.acceptProposedAction()
        self.setStyleSheet("QFrame { border: 2px solid #ff6b45; border-radius: 8px; background: #252a42; }")
    def dragLeaveEvent(self, event):
        self.setStyleSheet("QFrame { border: 2px solid #e8542e; border-radius: 8px; background: #1e2235; }")

    def dropEvent(self, event):
        self.setStyleSheet("QFrame { border: 2px solid #e8542e; border-radius: 8px; background: #1e2235; }")
        urls = event.mimeData().urls()
        if urls:
            path = urls[0].toLocalFile()
            if path:
                self.show_filename(path)
                self.file_dropped.emit(path)


class CueSheetPage(QWidget):
    back_requested = Signal()

    def __init__(self, parent=None):
        super().__init__(parent)
        self.parsed_data = None
        self.channel_checkboxes = []
        self.cues = []
        self._build_ui()

    def _build_ui(self):
        layout = QVBoxLayout(self)
        layout.setContentsMargins(24, 24, 24, 16)

        top_bar = QHBoxLayout()
        self._t_back = QPushButton(tr("back")); back_btn = self._t_back
        back_btn.clicked.connect(self.back_requested.emit)
        top_bar.addWidget(back_btn)
        top_bar.addStretch()
        layout.addLayout(top_bar)

        self._t_title = QLabel(tr("cs_title")); title = self._t_title
        title.setStyleSheet("font-size: 18px; font-weight: bold;")
        layout.addWidget(title)

        self.drop_zone = EdlDropZone()
        self.drop_zone.file_dropped.connect(self._load_file)
        layout.addWidget(self.drop_zone)

        browse_layout = QHBoxLayout()
        browse_btn = QPushButton("Browse...")
        browse_btn.clicked.connect(self._browse_file)
        browse_layout.addWidget(browse_btn)
        self.file_label = QLabel("No file loaded.")
        self.file_label.setStyleSheet("color: #555555;")
        browse_layout.addWidget(self.file_label)
        browse_layout.addStretch()
        layout.addLayout(browse_layout)

        settings_row = QHBoxLayout()

        fps_box = QVBoxLayout()
        fps_box.addWidget(TLabel("cs_fps"))
        self.fps_combo = QComboBox()
        self.fps_combo.addItems(list(FRAME_RATES.keys()))
        self.fps_combo.setCurrentText("25 (PAL)")
        fps_box.addWidget(self.fps_combo)
        settings_row.addLayout(fps_box)

        title_box = QVBoxLayout()
        title_box.addWidget(TLabel("doc_title"))
        self.document_title_input = QLineEdit()
        self.document_title_input.setPlaceholderText("Show / project name")
        title_box.addWidget(self.document_title_input)
        settings_row.addLayout(title_box)

        layout.addLayout(settings_row)

        self.channels_group = TGroupBox("cs_tracks")
        self.channels_layout = QHBoxLayout(self.channels_group)
        self.channels_group.setVisible(False)
        layout.addWidget(self.channels_group)

        generate_layout = QHBoxLayout()
        self.generate_btn = TButton("cs_generate")
        self.generate_btn.setMinimumHeight(34)
        self.generate_btn.clicked.connect(self._generate_table)
        self.generate_btn.setEnabled(False)
        generate_layout.addWidget(self.generate_btn)
        generate_layout.addStretch()
        layout.addLayout(generate_layout)

        self.table = QTableWidget()
        self.table.setColumnCount(len(COLUMNS))
        self.table.setHorizontalHeaderLabels(COLUMNS)
        self.table.horizontalHeader().setSectionResizeMode(1, QHeaderView.Stretch)
        self.table.setAlternatingRowColors(True)
        layout.addWidget(self.table)

        self.total_label = QLabel("")
        self.total_label.setStyleSheet("color: #555555;")
        layout.addWidget(self.total_label)

        export_group = TGroupBox("export_group")
        export_layout = QHBoxLayout(export_group)
        export_pdf_btn = TButton("cs_export_pdf")
        export_pdf_btn.clicked.connect(lambda: self._export("PDF"))
        export_layout.addWidget(export_pdf_btn)
        export_docx_btn = TButton("cs_export_word")
        export_docx_btn.clicked.connect(lambda: self._export("DOCX"))
        export_layout.addWidget(export_docx_btn)
        export_txt_btn = TButton("cs_export_txt")
        export_txt_btn.clicked.connect(lambda: self._export("TXT"))
        export_layout.addWidget(export_txt_btn)
        export_layout.addStretch()
        layout.addWidget(export_group)

    # ── Chargement du fichier ───────────────────────────────────────────

    def _browse_file(self):
        path, _ = QFileDialog.getOpenFileName(
            self, "Choisir un fichier EDL", "",
            "Fichiers EDL (*.edl *.EDL);;Tous les fichiers (*.*)"
        )
        if path:
            self._load_file(path)

    def _load_file(self, path):
        if not os.path.isfile(path):
            return
        if hasattr(self, "drop_zone"):
            self.drop_zone.show_filename(path)
        self.file_label.setText("")
        self.table.setRowCount(0)
        self.cues = []
        self.total_label.setText("")

        try:
            self.parsed_data = parse_edl(path)
        except EdlParseError as e:
            self.parsed_data = None
            self.channels_group.setVisible(False)
            self.generate_btn.setEnabled(False)
            play_error()
            QMessageBox.critical(self, "EDL read error", str(e))
            return
        except Exception as e:
            self.parsed_data = None
            self.channels_group.setVisible(False)
            self.generate_btn.setEnabled(False)
            play_error()
            QMessageBox.critical(self, "Unexpected error", str(e))
            return

        self.document_title_input.setText(self.parsed_data["title"])

        # Suggestion automatique du frame rate selon la mention drop-frame
        # trouvée dans l'EDL (simple indication, l'utilisateur peut corriger).
        if self.parsed_data.get("drop_frame_hint"):
            self.fps_combo.setCurrentText("29.97 (NTSC drop-frame)")

        self._populate_channel_checkboxes()
        self.generate_btn.setEnabled(True)

    def _populate_channel_checkboxes(self):
        while self.channels_layout.count():
            item = self.channels_layout.takeAt(0)
            widget = item.widget()
            if widget:
                widget.deleteLater()
        self.channel_checkboxes = []

        channels = sorted(set(e["channel"] for e in self.parsed_data["events"]))
        for ch in channels:
            count = sum(1 for e in self.parsed_data["events"] if e["channel"] == ch)
            checkbox = QCheckBox(f"Piste {ch} ({count} clip{'s' if count > 1 else ''})")
            checkbox.setChecked(True)
            checkbox.channel_id = ch
            self.channels_layout.addWidget(checkbox)
            self.channel_checkboxes.append(checkbox)

        self.channels_layout.addStretch()
        self.channels_group.setVisible(True)

    # ── Génération du tableau ───────────────────────────────────────────

    def _generate_table(self):
        if not self.parsed_data:
            return

        selected_channels = {cb.channel_id for cb in self.channel_checkboxes if cb.isChecked()}
        if not selected_channels:
            QMessageBox.warning(self, "No track selected",
                                "Check at least one audio track.")
            return

        fps_label = self.fps_combo.currentText()
        fps = FRAME_RATES[fps_label]
        drop_frame = "drop-frame" in fps_label.lower()

        cues = []
        for event in self.parsed_data["events"]:
            if event["channel"] not in selected_channels:
                continue
            try:
                frames_in = timecode_to_frames(event["rec_in"], fps, drop_frame)
                frames_out = timecode_to_frames(event["rec_out"], fps, drop_frame)
                duration_frames = max(0, frames_out - frames_in)
                duration_tc = frames_to_timecode(duration_frames, fps)
            except Exception:
                duration_tc = "?"

            cues.append({
                "event_num": event["event_num"],
                "name": event["name"],
                "channel": event["channel"],
                "tc_in": event["rec_in"],
                "tc_out": event["rec_out"],
                "duration": duration_tc,
                "duration_frames": duration_frames if isinstance(duration_frames, int) else 0,
            })

        cues.sort(key=lambda c: (c["tc_in"], c["channel"]))
        cues = self._merge_consecutive_cues(cues, fps, drop_frame)
        self.cues = cues
        self._refresh_table(fps)

    def _merge_consecutive_cues(self, cues, fps, drop_frame):
        """Fusionne les lignes consécutives qui correspondent à la même
        musique découpée par des coupes de montage.

        Deux lignes sont fusionnées si :
          - même nom de fichier/clip (name)
          - même piste (channel)
          - le TC-out de l'une touche le TC-in de la suivante à 1 image près

        IMPORTANT : la fusion se fait PISTE PAR PISTE. Une fois les cues triés
        par timecode, un événement d'une autre piste peut s'insérer
        chronologiquement entre deux segments d'une même musique (ex: musique
        sur A2 coupée, riser sur A au milieu). Si on comparait seulement les
        lignes voisines dans la liste globale, on raterait la fusion. On
        regroupe donc par (name, channel), on fusionne à l'intérieur, puis on
        recombine et on retrie."""
        if not cues:
            return cues

        from collections import OrderedDict

        groups = OrderedDict()
        for cue in cues:
            key = (cue["name"], cue["channel"])
            groups.setdefault(key, []).append(cue)

        merged_all = []
        for key, group in groups.items():
            def _frames(c):
                try:
                    return timecode_to_frames(c["tc_in"], fps, drop_frame)
                except Exception:
                    return 0
            group.sort(key=_frames)

            current = dict(group[0])
            for nxt in group[1:]:
                consecutive = False
                try:
                    cur_out = timecode_to_frames(current["tc_out"], fps, drop_frame)
                    nxt_in = timecode_to_frames(nxt["tc_in"], fps, drop_frame)
                    gap = nxt_in - cur_out
                    consecutive = (0 <= gap <= 1)
                except Exception:
                    consecutive = False

                if consecutive:
                    current["tc_out"] = nxt["tc_out"]
                    try:
                        f_in = timecode_to_frames(current["tc_in"], fps, drop_frame)
                        f_out = timecode_to_frames(current["tc_out"], fps, drop_frame)
                        dur_frames = max(0, f_out - f_in)
                        current["duration"] = frames_to_timecode(dur_frames, fps)
                        current["duration_frames"] = dur_frames
                    except Exception:
                        pass
                else:
                    merged_all.append(current)
                    current = dict(nxt)
            merged_all.append(current)

        merged_all.sort(key=lambda c: (c["tc_in"], c["channel"]))
        return merged_all

    def _refresh_table(self, fps):
        self.table.setRowCount(len(self.cues))
        for row, cue in enumerate(self.cues):
            values = [
                cue["event_num"], cue["name"], cue["channel"],
                cue["tc_in"], cue["tc_out"], cue["duration"]
            ]
            for col, value in enumerate(values):
                item = QTableWidgetItem(value)
                if col != 1:
                    item.setTextAlignment(Qt.AlignCenter)
                self.table.setItem(row, col, item)

        total_frames = sum(c["duration_frames"] for c in self.cues)
        total_tc = frames_to_timecode(total_frames, fps)
        self.total_label.setText(
            f"{len(self.cues)} track(s) — total music duration: {total_tc}"
        )

    # ── Export ───────────────────────────────────────────────────────

    def _export(self, fmt):
        if not self.cues:
            QMessageBox.warning(self, "Nothing to export", "Generate the table first.")
            return

        doc_title = self.document_title_input.text().strip() or "Cue Sheet"
        filters = {
            "PDF": "Document PDF (*.pdf)",
            "DOCX": "Document Word (*.docx)",
            "TXT": "Fichier texte (*.txt)",
        }
        from core.paths import default_output_dir as _dod
        path, _ = QFileDialog.getSaveFileName(self, "Exporter", os.path.join(_dod(), doc_title), filters[fmt])
        if not path:
            return

        try:
            if fmt == "PDF":
                self._export_pdf(path, doc_title)
            elif fmt == "DOCX":
                self._export_docx(path, doc_title)
            else:
                self._export_txt(path, doc_title)
            play_done()
            QMessageBox.information(self, "Export successful", f"File saved:\n{path}")
        except ImportError as e:
            play_error()
            QMessageBox.critical(self, "Missing dependency",
                f"{e}\nLance : pip install -r requirements.txt")
        except Exception as e:
            QMessageBox.critical(self, "Erreur d'export", str(e))

    def _export_txt(self, path, doc_title):
        col_widths = [6, 40, 8, 12, 12, 12]
        header = "  ".join(c.ljust(col_widths[i]) for i, c in enumerate(COLUMNS))
        separator = "-" * len(header)
        lines = [doc_title, "=" * len(doc_title), "", header, separator]
        for cue in self.cues:
            values = [cue["event_num"], cue["name"], cue["channel"],
                      cue["tc_in"], cue["tc_out"], cue["duration"]]
            lines.append("  ".join(str(v).ljust(col_widths[i]) for i, v in enumerate(values)))
        lines.extend(["", separator, self.total_label.text()])
        with open(path, "w", encoding="utf-8") as f:
            f.write("\n".join(lines))

    def _export_docx(self, path, doc_title):
        from docx import Document
        from docx.shared import Pt
        doc = Document()
        doc.add_heading(doc_title, level=1)
        table = doc.add_table(rows=1, cols=len(COLUMNS))
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        for i, col_name in enumerate(COLUMNS):
            hdr[i].text = col_name
        for cue in self.cues:
            row = table.add_row().cells
            row[0].text = cue["event_num"]
            row[1].text = cue["name"]
            row[2].text = cue["channel"]
            row[3].text = cue["tc_in"]
            row[4].text = cue["tc_out"]
            row[5].text = cue["duration"]
        doc.add_paragraph("")
        doc.add_paragraph(self.total_label.text())
        doc.save(path)

    def _export_pdf(self, path, doc_title):
        from reportlab.lib.pagesizes import A4, landscape
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.lib import colors

        doc = SimpleDocTemplate(path, pagesize=landscape(A4))
        styles = getSampleStyleSheet()
        elements = [Paragraph(doc_title, styles["Title"]), Spacer(1, 16)]

        data = [COLUMNS] + [
            [c["event_num"], c["name"], c["channel"],
             c["tc_in"], c["tc_out"], c["duration"]]
            for c in self.cues
        ]
        col_widths = [40, 220, 50, 85, 85, 85]
        table = Table(data, colWidths=col_widths, repeatRows=1)
        table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#333333")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 9),
            ("ALIGN", (0, 0), (-1, -1), "CENTER"),
            ("ALIGN", (1, 1), (1, -1), "LEFT"),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1),
             [colors.white, colors.HexColor("#f5f5f5")]),
        ]))
        elements.append(table)
        elements.append(Spacer(1, 16))
        elements.append(Paragraph(self.total_label.text(), styles["Normal"]))
        doc.build(elements)

    def set_lang(self, lang):
        from core.lang import tr, LangManager
        LangManager.get().current = lang
        if hasattr(self, '_t_title'):       self._t_title.setText(tr("cs_title"))
        if hasattr(self, '_t_back'):        self._t_back.setText(tr("back"))
        if hasattr(self, '_t_generate'):    self._t_generate.setText(tr("cs_generate"))
        if hasattr(self, '_t_exp_pdf'):     self._t_exp_pdf.setText(tr("cs_export_pdf"))
        if hasattr(self, '_t_exp_word'):    self._t_exp_word.setText(tr("cs_export_word"))
        if hasattr(self, '_t_exp_txt'):     self._t_exp_txt.setText(tr("cs_export_txt"))

